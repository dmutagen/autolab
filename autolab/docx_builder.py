import docx
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import re
from autolab.config import AppConfig

def normalize_dashes(text: str) -> str:
    """Replace long em-dashes (—) with middle en-dashes (–)."""
    if not text:
        return ""
    return text.replace("—", "–")

def clean_markdown(text: str) -> str:
    """Strip raw markdown artifacts like **bold**, `code`, and markdown headers."""
    if not text:
        return ""
    # Strip markdown code blocks
    text = re.sub(r'```[a-zA-Z]*\n?', '', text)
    text = text.replace('```', '')
    # Strip bold asterisks: **text** -> text
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Strip inline backticks: `code` -> code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Strip markdown headers: ### Header -> Header
    text = re.sub(r'^\s*#{1,6}\s+', '', text, flags=re.MULTILINE)
    return text.strip()

def lowercase_first_russian(word: str) -> str:
    """Lowercase first Cyrillic letter unless it is an all-caps acronym (ПЭВМ, ОС, СУБД, API, etc.)."""
    if not word:
        return ""
    if word.isupper() and len(word) >= 2:
        return word
    if word[0] in 'АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ':
        if len(word) == 1:
            return word.lower()
        if word[1].isupper():
            return word
        return word[0].lower() + word[1:]
    return word

def fix_colons_and_bullets(text: str) -> str:
    """
    Lowercase the first Russian word after colons (: or : - or : –) and list dashes (– ),
    as required by GOST and Russian typography (e.g. '4. Перезапуск Activity: для того чтобы...').
    Also normalizes spaces so there is exactly one space after ':' and '–'.
    """
    if not text:
        return ""

    # Normalize multiple spaces after bullet dash or numbers
    text = re.sub(r'(^[ \t]*[–—\-])[ \t]+', r'– ', text, flags=re.MULTILINE)
    text = re.sub(r'(^\d+[\.\)])[ \t]+', r'\1 ', text, flags=re.MULTILINE)

    # Normalize multiple spaces after colon
    text = re.sub(r':\s+', ': ', text)

    def repl_colon(m):
        prefix = m.group(1) # e.g. ': ', ': - ', ': – '
        word = m.group(2)
        return prefix + lowercase_first_russian(word)

    text = re.sub(r'(:[ \t]*(?:[–—\-][ \t]*)?)([A-Za-zА-Яа-яЁё]+)', repl_colon, text)

    def repl_bullet(m):
        prefix = m.group(1) # '– '
        word = m.group(2)
        return prefix + lowercase_first_russian(word)

    text = re.sub(r'(^[ \t]*[–—\-][ \t]+)([A-Za-zА-Яа-яЁё]+)', repl_bullet, text, flags=re.MULTILINE)
    return text

def sanitize_caption(caption: str) -> str:
    """
    Strip any existing 'Рисунок X – ', 'Рис. X: ', etc., to prevent duplicate
    'Рисунок 1 – Рисунок 1 – ...' captions.
    """
    if not caption:
        return "Результат выполнения работы"
    cap = clean_markdown(normalize_dashes(caption.strip(". ")))
    pattern = r'^(?:(?:рисунок|рис|иллюстрация|скриншот)\s*\d*[\s–—\-\.\:]*)+'
    while re.search(pattern, cap, flags=re.IGNORECASE):
        cap = re.sub(pattern, '', cap, flags=re.IGNORECASE).strip()
    cap = cap.lstrip("–—-:. ").strip()
    return cap if cap else "Результат выполнения работы"

def normalize_paragraphs(text: str) -> List[str]:
    """
    Split text into coherent paragraphs and list items, un-wrapping artificial
    line breaks so that Word JUSTIFY does not stretch words across the entire page.
    Replaces all list bullets (*, •, -, +) with middle en-dashes (–).
    Enforces lowercase after colons and list bullets.
    """
    if not text:
        return []

    clean_text = clean_markdown(normalize_dashes(text))
    raw_blocks = clean_text.split("\n\n")
    paragraphs = []

    for block in raw_blocks:
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        if not lines:
            continue

        current = []
        for line in lines:
            # Replace bullet markers (*, •, +, -, –, —) with middle en-dash (–)
            line = re.sub(r'^[ \t]*[\*•\-\–—\+][ \t]+', '– ', line)
            # Normalize spaces after numbers '1.  ' -> '1. '
            line = re.sub(r'^(\d+[\.\)])\s+', r'\1 ', line)
            # Enforce lowercase after colons and bullets
            line = fix_colons_and_bullets(line)

            # Check if line starts a list item: '1. ', '1) ', '– '
            is_list_item = bool(re.match(r'^(?:\d+[\.\)]|–)\s+', line))
            if is_list_item:
                if current:
                    paragraphs.append(" ".join(current))
                    current = []
                current.append(line)
            else:
                if current:
                    current.append(line)
                else:
                    current.append(line)

        if current:
            paragraphs.append(" ".join(current))

    return paragraphs

def get_short_filename(subject: str, lab_number: str) -> str:
    """Generate short filenames matching the student's catalog: МОБ_5.docx, разраб2.docx, СУБД_1.docx, etc."""
    num_match = re.search(r'\d+(?:[\.\-_]\d+)?', str(lab_number))
    num = num_match.group(0) if num_match else "1"
    
    s = (subject or "").lower().strip()
    if any(w in s for w in ["трпо"]):
        return f"ТРПО_{num}.docx"
    elif any(w in s for w in ["дизайн", "design", "ui", "ux", "figma", "макет"]):
        return f"дизайн_{num}.docx"
    elif any(w in s for w in ["мобил", "android"]):
        return f"МОБ_{num}.docx"
    elif any(w in s for w in ["субд", "баз", "данн", "sql", "sqlite", "access"]):
        return f"СУБД_{num}.docx"
    elif any(w in s for w in ["защит"]):
        return f"защитакомп{num}.docx"
    elif any(w in s for w in ["иб", "безопасн"]):
        return f"ИБ_{num}.docx"
    elif any(w in s for w in ["схем"]):
        return f"схема{num}.docx"
    elif any(w in s for w in ["микро", "stm", "avr", "пмк"]):
        return f"прогмикро{num}.docx"
    elif any(w in s for w in ["практик"]):
        return f"практика_{num}.docx"
    elif any(w in s for w in ["разраб", "поит", "java", "программ"]):
        return f"разраб{num}.docx"
    else:
        words = re.findall(r'[A-Za-zА-Яа-яЁё]+', s)
        prefix = words[0][:6] if words else "лаб"
        return f"{prefix}_{num}.docx"

class DocxBuilder:
    def __init__(self, config: AppConfig):
        self.config = config
        self.doc = docx.Document()
        self._setup_document_geometry()
        self._setup_default_styles()

    def _setup_document_geometry(self):
        section = self.doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        # Belarusian GOST: Left 30 mm, Right 10 mm, Top 20 mm, Bottom 20 mm
        section.left_margin = Mm(self.config.gost.margin_left_mm)
        section.right_margin = Mm(self.config.gost.margin_right_mm)
        section.top_margin = Mm(self.config.gost.margin_top_mm)
        section.bottom_margin = Mm(self.config.gost.margin_bottom_mm)

    def _setup_default_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0, 0, 0)
        p_format = style.paragraph_format
        p_format.line_spacing = 1.0
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.first_line_indent = Mm(self.config.gost.first_line_indent_cm * 10)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_paragraph(
        self,
        text: str = "",
        bold: bool = False,
        italic: bool = False,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
        indent_cm: Optional[float] = None,
        font_size_pt: Optional[float] = None
    ):
        p = self.doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        if indent_cm is not None:
            p.paragraph_format.first_line_indent = Mm(indent_cm * 10)
        else:
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                p.paragraph_format.first_line_indent = Mm(0)
            else:
                p.paragraph_format.first_line_indent = Mm(self.config.gost.first_line_indent_cm * 10)

        if text:
            clean_text = clean_markdown(normalize_dashes(text))
            # Protect official metadata lines from being lowercased
            is_metadata = any(clean_text.startswith(prefix) for prefix in [
                "Номер учебной группы:", "Фамилия, инициалы", "Дата выполнения", "Тема работы:",
                "Оснащение работы:", "Вариант:", "Выполнил:", "Принял:", "Вывод:", "Рисунок"
            ])
            if not is_metadata:
                clean_text = fix_colons_and_bullets(clean_text)

            run = p.add_run(clean_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size_pt or 14)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        return p

    def _add_text_block(self, text: str, bold: bool = False, prefix: str = ""):
        """Add multi-paragraph text cleanly without stretching line breaks."""
        paras = normalize_paragraphs(text)
        for i, p_text in enumerate(paras):
            full_text = f"{prefix}{p_text}" if i == 0 and prefix else p_text
            self._add_paragraph(full_text, bold=bold)

    def add_header(self, data: Dict[str, Any], user_variant: str = ""):
        lab_type = data.get("lab_type", "Лабораторная работа")
        lab_num = data.get("lab_number", "1")
        student = self.config.student
        
        # 1. Title: regular font, centered
        self._add_paragraph(
            f"{lab_type} № {lab_num}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            indent_cm=0
        )

        # 2. Metadata lines: regular font, 1.25 cm indent
        self._add_paragraph(f"Номер учебной группы: {student.group}")
        self._add_paragraph(f"Фамилия, инициалы обучающегося: {student.student_name}")
        
        raw_date = data.get("date") or ""
        if raw_date:
            if re.match(r'^\d{4}-\d{2}-\d{2}$', raw_date):
                parts = raw_date.split('-')
                date_str = f"{parts[2]}.{parts[1]}.{parts[0]}"
            else:
                date_str = raw_date
        else:
            date_str = datetime.now().strftime("%d.%m.%Y")
            
        self._add_paragraph(f"Дата выполнения работы: {date_str}")

        topic = clean_markdown(data.get("topic", "").strip("«»\" "))
        self._add_paragraph(f"Тема работы: «{topic}»")

        goal = clean_markdown(data.get("goal", "").strip())
        goal = re.sub(r'^\s*цель(?:\s+работы)?\s*:\s*', '', goal, flags=re.IGNORECASE)
        if goal:
            self._add_text_block(goal, prefix="Цель работы: ")

        task = clean_markdown(data.get("task", "").strip())
        task = re.sub(r'^\s*задание\s*:\s*', '', task, flags=re.IGNORECASE)
        if task:
            self._add_text_block(task, prefix="Задание: ")

        # ONLY add variant if user entered a variant explicitly on the site!
        clean_var = user_variant.strip() if user_variant else ""
        if clean_var and clean_var.lower() not in ["общий", "согласно заданию", "нет", "none", "-"]:
            self._add_paragraph(f"Вариант: {clean_var}")

        equip = clean_markdown(data.get("equipment", "").strip())
        equip = re.sub(r'^\s*оснащение(?:\s+работы)?\s*:\s*', '', equip, flags=re.IGNORECASE)
        if equip:
            self._add_text_block(equip, prefix="Оснащение работы: ")

    def add_code_block(self, code: str):
        self._add_paragraph("Код программы:")
        clean_code = code.strip()
        clean_code = re.sub(r'^```[a-zA-Z]*\n', '', clean_code)
        clean_code = re.sub(r'\n```$', '', clean_code)
        
        lines = clean_code.split("\n")
        for line in lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Mm(0)
            
            run = p.add_run(line if line else " ")
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    def add_screenshot(self, image_path: Path, figure_number: int, caption: str):
        if not image_path.exists():
            return

        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.line_spacing = 1.0
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        p_img.paragraph_format.first_line_indent = Mm(0)
        
        run_img = p_img.add_run()
        run_img.add_picture(str(image_path), width=Mm(160))

        clean_cap = sanitize_caption(caption)
        self._add_paragraph(
            f"Рисунок {figure_number} – {clean_cap}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            indent_cm=0,
            font_size_pt=14
        )

    def _remove_table_borders(self, table):
        tblPr = table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    def _set_cell_padding(self, cell, top=0, bottom=40, left=15, right=15):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_screenshots_side_by_side(self, screenshots: List[Path], start_fig_no: int, captions: List[str]):
        """
        Embeds multiple screenshots horizontally on the same line in a single borderless table.
        Each column has an image and its corresponding caption 'Рисунок X – <caption_i>'.
        """
        valid_pairs = []
        for i, s in enumerate(screenshots):
            if s.exists():
                cap = captions[i] if i < len(captions) else "Результат выполнения программы"
                valid_pairs.append((s, sanitize_caption(cap), start_fig_no + i))

        n = len(valid_pairs)
        if n == 0:
            return
        if n == 1:
            s, cap, f_no = valid_pairs[0]
            self.add_screenshot(s, f_no, cap)
            return

        table = self.doc.add_table(rows=2, cols=n)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._remove_table_borders(table)

        col_w_mm = 168.0 / n
        img_w_mm = max(col_w_mm - 2, 15.0)

        for i, (shot_path, cap_text, fig_num) in enumerate(valid_pairs):
            for row in table.rows:
                row.cells[i].width = Mm(col_w_mm)
                self._set_cell_padding(row.cells[i], top=0, bottom=40, left=15, right=15)

            # Row 0: Image
            c_img = table.cell(0, i)
            p_img = c_img.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.line_spacing = 1.0
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(0)
            p_img.paragraph_format.first_line_indent = Mm(0)
            r_img = p_img.add_run()
            r_img.add_picture(str(shot_path), width=Mm(img_w_mm))

            # Row 1: Caption
            c_cap = table.cell(1, i)
            p_cap = c_cap.paragraphs[0]
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.line_spacing = 1.0
            p_cap.paragraph_format.space_before = Pt(4)
            p_cap.paragraph_format.space_after = Pt(0)
            p_cap.paragraph_format.first_line_indent = Mm(0)
            r_cap = p_cap.add_run(f"Рисунок {fig_num} – {cap_text}")
            r_cap.font.name = "Times New Roman"
            r_cap.font.size = Pt(12 if n >= 4 else (13 if n == 3 else 14))

        p_after = self.doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(0)
        p_after.paragraph_format.space_after = Pt(6)
        p_after.paragraph_format.line_spacing = 1.0
        p_after.paragraph_format.first_line_indent = Mm(0)

    def add_title_page(self, data: Dict[str, Any]):
        student = self.config.student
        
        self._add_paragraph("МИНИСТЕРСТВО ОБРАЗОВАНИЯ РЕСПУБЛИКИ БЕЛАРУСЬ", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)
        self._add_paragraph(student.institution.upper(), alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)
        self._add_paragraph(f"Специальность {student.specialty}", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)
        self._add_paragraph(f"Группа {student.group}", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)

        for _ in range(3):
            self._add_paragraph("", indent_cm=0)

        lab_num = data.get("lab_number", "1")
        self._add_paragraph("ОТЧЕТ", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0, font_size_pt=16)
        self._add_paragraph(f"по лабораторной работе № {lab_num}", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)
        
        topic = data.get("topic", "")
        if topic:
            self._add_paragraph(f"«{clean_markdown(topic)}»", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)

        for _ in range(4):
            self._add_paragraph("", indent_cm=0)

        self._add_paragraph(f"Выполнил: обучающийся группы {student.group} {student.student_name}", indent_cm=8)
        self._add_paragraph(f"Принял: преподаватель {student.teacher_name}", indent_cm=8)

        for _ in range(4):
            self._add_paragraph("", indent_cm=0)

        self._add_paragraph(f"{student.city}  {student.year}", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent_cm=0)
        self.doc.add_page_break()

    def build_report(
        self,
        data: Dict[str, Any],
        output_file: Path,
        screenshots: List[Path],
        with_title_page: bool = False,
        include_theory: bool = False,
        user_variant: str = "",
        photos_layout: str = "side_by_side",
        include_screenshot_intro: bool = True
    ) -> Path:
        if with_title_page:
            self.add_title_page(data)

        self.add_header(data, user_variant=user_variant)

        # Section: Результаты выполнения работы
        self._add_paragraph("Результаты выполнения работы:")

        code = data.get("code", "").strip()
        sol_desc = data.get("solution_description", "").strip()

        if code:
            # Coding lab:
            if include_theory:
                theory = data.get("theory", "").strip()
                if theory:
                    self._add_text_block(theory)
                if sol_desc:
                    self._add_text_block(sol_desc)
            self.add_code_block(code)
        else:
            # Design / Non-coding lab (Figma, UI/UX, Schemes, Modeling):
            if include_theory:
                theory = data.get("theory", "").strip()
                if theory:
                    self._add_text_block(theory)
            if sol_desc:
                self._add_text_block(sol_desc)

        # Screenshots (can be multiple, side-by-side or separate)
        if screenshots:
            if include_screenshot_intro:
                n_shots = len(screenshots)
                if n_shots == 1:
                    intro_text = "Внешний вид разработанного интерфейса и результат работы программы представлены на рисунке 1:"
                else:
                    intro_text = f"Внешний вид разработанного интерфейса и результаты выполнения работы представлены на рисунках 1–{n_shots}:"
                self._add_paragraph(intro_text)

            figures_meta = data.get("figures", [])
            if photos_layout == "side_by_side" and len(screenshots) > 1:
                caps = []
                for i in range(len(screenshots)):
                    c = "Результат выполнения программы" if code else "Макет разработанного интерфейса"
                    if i < len(figures_meta) and "title" in figures_meta[i]:
                        c = figures_meta[i]["title"]
                    elif len(screenshots) > 1:
                        c = f"Окно программы (часть {i+1})" if code else f"Экран интерфейса ({i+1})"
                    caps.append(c)
                self.add_screenshots_side_by_side(screenshots, 1, caps)
            else:
                fig_no = 1
                for i, shot in enumerate(screenshots):
                    cap = "Макет разработанного интерфейса" if not code else "Результат выполнения программы"
                    if i < len(figures_meta) and "title" in figures_meta[i]:
                        cap = figures_meta[i]["title"]
                    elif len(screenshots) > 1:
                        cap = f"Интерфейс программы (часть {i+1})" if code else f"Макет интерфейса (экран {i+1})"
                    self.add_screenshot(shot, fig_no, cap)
                    fig_no += 1

        # Control Questions
        qa_list = data.get("questions_answers", [])
        if qa_list:
            self._add_paragraph("Ответы на контрольные вопросы:")
            for i, qa in enumerate(qa_list):
                raw_q = clean_markdown(qa.get("question", "").strip())
                clean_q = re.sub(r'^\s*(\d+[\.\)]\s*)+', '', raw_q).strip()
                a = clean_markdown(qa.get("answer", "").strip())
                
                self._add_paragraph(f"{i+1}. {clean_q}")
                self._add_text_block(a)

        # Conclusion: Вывод
        concl = clean_markdown(data.get("conclusion", "").strip())
        if concl:
            if concl.lower().startswith("вывод:"):
                concl_text = concl[6:].strip()
            else:
                concl_text = concl
            self._add_paragraph(f"Вывод: {concl_text}")

        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_file))
        return output_file
