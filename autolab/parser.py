import os
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from pypdf import PdfReader
import docx
from PIL import Image

class TaskParser:
    @staticmethod
    def parse_docx(path: Path) -> str:
        try:
            doc = docx.Document(path)
            full_text = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    full_text.append(t)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if row_text:
                        full_text.append(row_text)
            return "\n".join(full_text)
        except Exception as e:
            return f"Ошибка чтения DOCX: {e}"

    @staticmethod
    def parse_pdf(path: Path) -> Tuple[str, List[Path]]:
        text_content = []
        image_paths = []
        try:
            reader = PdfReader(str(path))
            for i, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if txt.strip():
                    text_content.append(f"--- Страница {i+1} ---\n{txt.strip()}")
            
            full_text = "\n\n".join(text_content)
            return full_text, image_paths
        except Exception as e:
            return f"Ошибка чтения PDF: {e}", image_paths

    @staticmethod
    def process_input(
        text_input: str = "",
        uploaded_files: Optional[List[Path]] = None,
        uploaded_file: Optional[Path] = None,
        variant: Optional[str] = None
    ) -> Dict[str, Any]:
        result = {
            "text": text_input.strip(),
            "images": [],
            "source_type": "text",
            "variant": variant or ""
        }

        all_files: List[Path] = []
        if uploaded_files:
            all_files.extend(uploaded_files)
        if uploaded_file and uploaded_file not in all_files:
            all_files.append(uploaded_file)

        extracted_texts = []
        for fpath in all_files:
            if not fpath or not fpath.exists():
                continue

            suffix = fpath.suffix.lower()
            if suffix in [".docx", ".doc"]:
                result["source_type"] = "docx"
                extracted = TaskParser.parse_docx(fpath)
                extracted_texts.append(f"=== МАТЕРИАЛ ИЗ ФАЙЛА ({fpath.name}) ===\n{extracted}")
            elif suffix == ".pdf":
                result["source_type"] = "pdf"
                extracted, images = TaskParser.parse_pdf(fpath)
                extracted_texts.append(f"=== МАТЕРИАЛ ИЗ МЕТОДИЧКИ ({fpath.name}) ===\n{extracted}")
                result["images"].extend(images)
            elif suffix in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
                result["source_type"] = "image"
                result["images"].append(fpath)
            elif suffix in [".txt", ".md"]:
                try:
                    with open(fpath, "r", encoding="utf-8", errors="replace") as tf:
                        extracted_texts.append(f"=== МАТЕРИАЛ ({fpath.name}) ===\n{tf.read()}")
                except Exception:
                    pass

        if extracted_texts:
            materials_str = "\n\n".join(extracted_texts)
            if result["text"]:
                result["text"] = f"{result['text']}\n\n{materials_str}"
            else:
                result["text"] = materials_str

        return result
